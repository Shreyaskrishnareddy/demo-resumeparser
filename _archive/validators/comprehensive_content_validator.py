#!/usr/bin/env python3
"""
Comprehensive Resume Content Validator
Validates extracted data against actual resume content with detailed analysis
"""

import os
import re
import time
import fitz
from docx import Document
from fast_brd_transformer import FastBRDTransformer

class ResumeContentValidator:
    def __init__(self):
        self.transformer = FastBRDTransformer()
        self.test_dir = "/home/great/claudeprojects/parser/test_resumes/Test Resumes"

    def get_file_text(self, file_path):
        """Extract text from various file formats"""
        file_path = str(file_path)
        extension = os.path.splitext(file_path)[1].lower()

        try:
            if extension == '.pdf':
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            elif extension == '.docx':
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            elif extension == '.doc':
                try:
                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    if text.strip():
                        return text
                except:
                    pass

                # Fallback: binary extraction
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read()
                        text_chunks = []
                        current_chunk = ""

                        for byte in content:
                            if 32 <= byte <= 126:  # Printable ASCII
                                current_chunk += chr(byte)
                            elif byte in [10, 13]:  # Newlines
                                if current_chunk:
                                    current_chunk += '\n'
                            else:
                                if len(current_chunk) >= 3:
                                    clean_chunk = current_chunk.strip()
                                    if (len(clean_chunk) >= 3 and
                                        sum(c.isalpha() for c in clean_chunk) > len(clean_chunk) * 0.3):
                                        text_chunks.append(clean_chunk)
                                current_chunk = ""

                        if len(current_chunk) >= 3:
                            clean_chunk = current_chunk.strip()
                            if (len(clean_chunk) >= 3 and
                                sum(c.isalpha() for c in clean_chunk) > len(clean_chunk) * 0.3):
                                text_chunks.append(clean_chunk)

                        text = '\n'.join(text_chunks)
                        text = re.sub(r'[^\w\s\.\@\-\(\)\,\;\:\!\?\#\$\%\&\*\+\=\[\]]+', ' ', text)
                        text = re.sub(r'\s+', ' ', text)
                        text = re.sub(r'\n\s*\n+', '\n', text)
                        return text
                except Exception as e:
                    print(f"Error reading .doc file: {e}")
                    return ""
            else:
                return ""
        except:
            return ""

    def analyze_resume_content(self, text):
        """Analyze the actual content of a resume to understand what should be extracted"""
        analysis = {
            'names_found': [],
            'emails_found': [],
            'phones_found': [],
            'companies_found': [],
            'job_titles_found': [],
            'dates_found': [],
            'education_found': [],
            'skills_found': [],
            'certifications_found': [],
            'locations_found': []
        }

        # Find all potential names (first and last name combinations)
        name_patterns = [
            r'\b([A-Z][a-z]{2,15}\s+[A-Z][a-z]{2,15})\b',  # First Last
            r'\b([A-Z][a-z]{2,15}\s+[A-Z]\.\s+[A-Z][a-z]{2,15})\b',  # First M. Last
            r'\b([A-Z][a-z]{2,15}\s+[A-Z][a-z]{2,15}\s+[A-Z][a-z]{2,15})\b'  # First Middle Last
        ]

        for pattern in name_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                # Filter out obvious non-names
                if not any(word in match.lower() for word in [
                    'university', 'college', 'company', 'technologies', 'solutions', 'services',
                    'software', 'development', 'management', 'project', 'business', 'data'
                ]):
                    analysis['names_found'].append(match)

        # Find emails
        email_pattern = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
        analysis['emails_found'] = re.findall(email_pattern, text)

        # Find phone numbers
        phone_patterns = [
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        ]
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            analysis['phones_found'].extend(matches)

        # Find companies
        company_patterns = [
            r'\b([A-Z][a-zA-Z\s&\.,]{5,40}?)\s+(Inc\.?|Ltd\.?|Corp\.?|LLC|Consulting|Technologies|Solutions|Services|Systems|Group)\b',
            r'\b([A-Z][a-zA-Z\s&\.,]{8,50})\s*(?=\s*[-–|]\s*[A-Z]|\s*\(|\s*,|\s*\n)',
        ]
        for pattern in company_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    company = ' '.join(match).strip()
                else:
                    company = match.strip()
                if len(company) > 5:
                    analysis['companies_found'].append(company)

        # Find job titles
        title_patterns = [
            r'\b(Senior|Junior|Lead|Principal|Staff|Associate)?\s*(Manager|Developer|Engineer|Analyst|Consultant|Director|Coordinator|Specialist|Administrator|Architect)\b',
            r'\b(Software|Data|Business|Product|Project|Program|Marketing|Sales|HR|Finance|System|Network|Database)\s+(Manager|Developer|Engineer|Analyst|Consultant|Administrator|Architect)\b',
        ]
        for pattern in title_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    title = ' '.join(match).strip()
                else:
                    title = match.strip()
                if title:
                    analysis['job_titles_found'].append(title)

        # Find dates
        date_patterns = [
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b',
            r'\b\d{1,2}/\d{4}\b',
            r'\b\d{4}\s*[-–]\s*(?:\d{4}|present|current)\b',
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b'
        ]
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            analysis['dates_found'].extend(matches)

        # Find education
        education_patterns = [
            r'\b(Bachelor|Master|PhD|Doctorate|Associate|Certificate|Diploma)\s+(?:of\s+)?(?:Science|Arts|Engineering|Business|Computer|Information)\b',
            r'\b(BS|BA|MS|MA|PhD|MBA|BE|BTech|MTech)\s+(?:in\s+)?[A-Za-z\s]{5,30}\b',
            r'\b([A-Z][a-zA-Z\s]{10,50})\s+(University|College|Institute|School)\b',
        ]
        for pattern in education_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    edu = ' '.join(match).strip()
                else:
                    edu = match.strip()
                analysis['education_found'].append(edu)

        # Find skills (technical terms)
        skill_keywords = [
            'Python', 'Java', 'JavaScript', 'React', 'Angular', 'Node.js', 'SQL', 'AWS', 'Docker',
            'Git', 'TypeScript', 'C++', 'C#', 'Azure', 'Machine Learning', 'Agile', 'Scrum', 'JIRA',
            'HTML', 'CSS', 'REST', 'API', 'MongoDB', 'PostgreSQL', 'MySQL', 'Redis', 'Kafka',
            'Kubernetes', 'Jenkins', 'Linux', 'Windows', 'Oracle', 'Tableau', 'Power BI'
        ]

        text_lower = text.lower()
        for skill in skill_keywords:
            if skill.lower() in text_lower:
                analysis['skills_found'].append(skill)

        # Find certifications
        cert_patterns = [
            r'\b(AWS|Microsoft|Google|Oracle|Cisco|PMI|CompTIA|CISSP|PMP|Scrum Master|CSM|CISA|CISM)\s+(?:Certified|Certification|Certificate)\b',
            r'\b(?:Certified|Certification)\s+([A-Za-z\s]{5,30})\b'
        ]
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    cert = ' '.join(match).strip()
                else:
                    cert = match.strip()
                analysis['certifications_found'].append(cert)

        return analysis

    def validate_extracted_data(self, extracted_data, content_analysis, original_text):
        """Validate extracted data against actual resume content"""
        validation_results = {
            'name_accuracy': 0,
            'email_accuracy': 0,
            'phone_accuracy': 0,
            'company_accuracy': 0,
            'job_title_accuracy': 0,
            'education_accuracy': 0,
            'skills_accuracy': 0,
            'date_accuracy': 0,
            'issues_found': [],
            'missing_data': [],
            'incorrect_data': [],
            'extraction_quality': 'Poor'
        }

        personal = extracted_data.get('PersonalDetails', {})
        experiences = extracted_data.get('ListOfExperiences', [])
        skills = extracted_data.get('ListOfSkills', [])
        education = extracted_data.get('Education', [])

        # Validate Name
        extracted_name = personal.get('FullName', '')
        if extracted_name:
            # Check if extracted name appears in the resume
            if any(name in original_text for name in content_analysis['names_found'] if extracted_name.lower() in name.lower()):
                validation_results['name_accuracy'] = 100
            elif extracted_name.lower() in original_text.lower():
                validation_results['name_accuracy'] = 80
            else:
                validation_results['name_accuracy'] = 0
                validation_results['incorrect_data'].append(f"Name '{extracted_name}' not found in resume")
        else:
            validation_results['missing_data'].append("No name extracted")

        # Validate Email
        extracted_email = personal.get('EmailID', '')
        if extracted_email:
            if extracted_email in content_analysis['emails_found']:
                validation_results['email_accuracy'] = 100
            elif any(email in extracted_email for email in content_analysis['emails_found']):
                validation_results['email_accuracy'] = 80
            else:
                validation_results['email_accuracy'] = 0
                validation_results['incorrect_data'].append(f"Email '{extracted_email}' not found in resume")
        else:
            validation_results['missing_data'].append("No email extracted")

        # Validate Phone
        extracted_phone = personal.get('PhoneNumber', '')
        if extracted_phone:
            # Clean phone for comparison
            clean_extracted = re.sub(r'[^\d]', '', extracted_phone)
            found_match = False
            for phone in content_analysis['phones_found']:
                clean_resume = re.sub(r'[^\d]', '', phone)
                if clean_extracted in clean_resume or clean_resume in clean_extracted:
                    validation_results['phone_accuracy'] = 100
                    found_match = True
                    break
            if not found_match:
                validation_results['phone_accuracy'] = 0
                validation_results['incorrect_data'].append(f"Phone '{extracted_phone}' not found in resume")
        else:
            validation_results['missing_data'].append("No phone extracted")

        # Validate Companies
        extracted_companies = [exp.get('CompanyName', '') for exp in experiences if exp.get('CompanyName')]
        if extracted_companies:
            matches = 0
            for ext_company in extracted_companies:
                for resume_company in content_analysis['companies_found']:
                    # Check partial matches
                    if (ext_company.lower() in resume_company.lower() or
                        resume_company.lower() in ext_company.lower() or
                        any(word in resume_company.lower() for word in ext_company.lower().split() if len(word) > 3)):
                        matches += 1
                        break
            validation_results['company_accuracy'] = min(100, (matches / len(extracted_companies)) * 100)
        else:
            validation_results['missing_data'].append("No companies extracted")

        # Validate Job Titles
        extracted_titles = [exp.get('JobTitle', '') for exp in experiences if exp.get('JobTitle')]
        if extracted_titles:
            matches = 0
            for ext_title in extracted_titles:
                for resume_title in content_analysis['job_titles_found']:
                    if (ext_title.lower() in resume_title.lower() or
                        resume_title.lower() in ext_title.lower()):
                        matches += 1
                        break
            validation_results['job_title_accuracy'] = min(100, (matches / len(extracted_titles)) * 100)
        else:
            validation_results['missing_data'].append("No job titles extracted")

        # Validate Education
        extracted_education = [edu.get('FullEducationDetails', '') for edu in education if edu.get('FullEducationDetails')]
        if extracted_education:
            matches = 0
            for ext_edu in extracted_education:
                for resume_edu in content_analysis['education_found']:
                    if (ext_edu.lower() in resume_edu.lower() or
                        resume_edu.lower() in ext_edu.lower() or
                        any(word in resume_edu.lower() for word in ext_edu.lower().split() if len(word) > 4)):
                        matches += 1
                        break
            validation_results['education_accuracy'] = min(100, (matches / len(extracted_education)) * 100)
        else:
            validation_results['missing_data'].append("No education extracted")

        # Validate Skills
        extracted_skills = [skill.get('SkillsName', '') for skill in skills if skill.get('SkillsName')]
        if extracted_skills:
            matches = 0
            for ext_skill in extracted_skills:
                if ext_skill in content_analysis['skills_found']:
                    matches += 1
            validation_results['skills_accuracy'] = min(100, (matches / len(extracted_skills)) * 100) if extracted_skills else 0
        else:
            validation_results['missing_data'].append("No skills extracted")

        # Calculate overall quality
        accuracies = [
            validation_results['name_accuracy'],
            validation_results['email_accuracy'],
            validation_results['phone_accuracy'],
            validation_results['company_accuracy'],
            validation_results['job_title_accuracy'],
            validation_results['education_accuracy'],
            validation_results['skills_accuracy']
        ]

        overall_accuracy = sum(accuracies) / len(accuracies)

        if overall_accuracy >= 90:
            validation_results['extraction_quality'] = 'Excellent'
        elif overall_accuracy >= 75:
            validation_results['extraction_quality'] = 'Good'
        elif overall_accuracy >= 60:
            validation_results['extraction_quality'] = 'Fair'
        else:
            validation_results['extraction_quality'] = 'Poor'

        return validation_results

    def comprehensive_validate_file(self, filename):
        """Comprehensive validation of a single resume file"""
        file_path = os.path.join(self.test_dir, filename)
        if not os.path.exists(file_path):
            return None

        print(f"\n🔍 COMPREHENSIVE VALIDATION: {filename}")
        print("=" * 80)

        # Extract text
        text = self.get_file_text(file_path)
        if not text:
            print("❌ Could not extract text")
            return None

        # Analyze actual content
        print("📄 ANALYZING ACTUAL RESUME CONTENT...")
        content_analysis = self.analyze_resume_content(text)

        print(f"  Names found in resume: {content_analysis['names_found'][:3]}")
        print(f"  Emails found: {content_analysis['emails_found']}")
        print(f"  Phones found: {content_analysis['phones_found']}")
        print(f"  Companies found: {content_analysis['companies_found'][:3]}")
        print(f"  Job titles found: {content_analysis['job_titles_found'][:3]}")
        print(f"  Education found: {content_analysis['education_found'][:2]}")
        print(f"  Skills found: {len(content_analysis['skills_found'])} skills")

        # Parse with our transformer
        print("\n🤖 PARSING WITH OUR TRANSFORMER...")
        start_time = time.time()
        extracted_data = self.transformer.transform_to_brd_format(text, filename)
        parse_time = (time.time() - start_time) * 1000

        # Display extracted data
        personal = extracted_data.get('PersonalDetails', {})
        experiences = extracted_data.get('ListOfExperiences', [])
        skills = extracted_data.get('ListOfSkills', [])
        education = extracted_data.get('Education', [])

        print(f"  ⏱️  Parse time: {parse_time:.2f}ms")
        print(f"  👤 Extracted name: '{personal.get('FullName', 'NOT FOUND')}'")
        print(f"  📧 Extracted email: '{personal.get('EmailID', 'NOT FOUND')}'")
        print(f"  📞 Extracted phone: '{personal.get('PhoneNumber', 'NOT FOUND')}'")
        print(f"  🏢 Extracted companies: {[exp.get('CompanyName', '') for exp in experiences[:3]]}")
        print(f"  💼 Extracted job titles: {[exp.get('JobTitle', '') for exp in experiences[:3]]}")
        print(f"  🎓 Extracted education: {[edu.get('FullEducationDetails', '') for edu in education[:2]]}")
        print(f"  🛠️  Extracted skills: {len(skills)} skills")

        # Validate extracted data against content
        print("\n✅ VALIDATING EXTRACTION ACCURACY...")
        validation = self.validate_extracted_data(extracted_data, content_analysis, text)

        print(f"  📊 Name accuracy: {validation['name_accuracy']:.0f}%")
        print(f"  📊 Email accuracy: {validation['email_accuracy']:.0f}%")
        print(f"  📊 Phone accuracy: {validation['phone_accuracy']:.0f}%")
        print(f"  📊 Company accuracy: {validation['company_accuracy']:.0f}%")
        print(f"  📊 Job title accuracy: {validation['job_title_accuracy']:.0f}%")
        print(f"  📊 Education accuracy: {validation['education_accuracy']:.0f}%")
        print(f"  📊 Skills accuracy: {validation['skills_accuracy']:.0f}%")
        print(f"  🎯 Overall quality: {validation['extraction_quality']}")

        # Report issues
        if validation['missing_data']:
            print(f"\n⚠️  MISSING DATA:")
            for missing in validation['missing_data']:
                print(f"    - {missing}")

        if validation['incorrect_data']:
            print(f"\n❌ INCORRECT DATA:")
            for incorrect in validation['incorrect_data']:
                print(f"    - {incorrect}")

        return {
            'filename': filename,
            'parse_time': parse_time,
            'content_analysis': content_analysis,
            'extracted_data': extracted_data,
            'validation': validation,
            'overall_accuracy': sum([
                validation['name_accuracy'],
                validation['email_accuracy'],
                validation['phone_accuracy'],
                validation['company_accuracy'],
                validation['job_title_accuracy'],
                validation['education_accuracy'],
                validation['skills_accuracy']
            ]) / 7
        }

    def validate_all_resumes(self):
        """Validate all target resume files"""
        target_files = [
            "Ahmad Qasem-Resume.pdf",
            "Ashok Kumar.doc",
            "Dexter Nigel Ramkissoon.docx",
            "Donald Belvin.docx",
            "Kiran N. Penmetcha_s Profile.pdf",
            "Mahesh_Bolikonda (1).pdf",
            "PRANAY REDDY_DE_Resume.pdf",
            "Resume of Connal Jackson.doc"
        ]

        print("🎯 COMPREHENSIVE RESUME CONTENT VALIDATION")
        print("=" * 80)

        results = []
        for filename in target_files:
            result = self.comprehensive_validate_file(filename)
            if result:
                results.append(result)

        # Summary
        print("\n" + "=" * 80)
        print("📊 VALIDATION SUMMARY:")

        if results:
            avg_accuracy = sum(r['overall_accuracy'] for r in results) / len(results)
            excellent_count = sum(1 for r in results if r['validation']['extraction_quality'] == 'Excellent')
            good_count = sum(1 for r in results if r['validation']['extraction_quality'] == 'Good')

            print(f"  Files validated: {len(results)}")
            print(f"  Average accuracy: {avg_accuracy:.1f}%")
            print(f"  Excellent quality: {excellent_count}/{len(results)}")
            print(f"  Good+ quality: {good_count + excellent_count}/{len(results)}")

            print(f"\n🎯 TARGET STATUS:")
            if avg_accuracy >= 90:
                print("  ✅ EXCEPTIONAL - Exceeding accuracy targets!")
            elif avg_accuracy >= 80:
                print("  ✅ EXCELLENT - Meeting high accuracy standards!")
            elif avg_accuracy >= 70:
                print("  ⚠️  GOOD - Close to target, minor improvements needed")
            else:
                print("  ❌ NEEDS IMPROVEMENT - Significant accuracy issues found")

        return results

if __name__ == "__main__":
    validator = ResumeContentValidator()
    validator.validate_all_resumes()