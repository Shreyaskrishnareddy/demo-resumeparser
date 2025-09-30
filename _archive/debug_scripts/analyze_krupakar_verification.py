#!/usr/bin/env python3
"""
Script to analyze Krupakar Reddy's resume verification data
"""

import pandas as pd
from docx import Document
import json
import sys
import os

def extract_docx_content(file_path):
    """Extract text content from DOCX file"""
    try:
        doc = Document(file_path)
        content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())

        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    content.append(" | ".join(row_data))

        return "\n".join(content)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def extract_excel_content(file_path):
    """Extract verification data from Excel file"""
    try:
        # Try reading all sheets
        xl_file = pd.ExcelFile(file_path)
        all_data = {}

        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            all_data[sheet_name] = df.to_dict('records')

        return all_data
    except Exception as e:
        print(f"Error reading Excel: {e}")
        return None

def analyze_verification_data():
    """Main analysis function"""

    # File paths
    excel_path = "/home/great/claudeprojects/parser/parserdemo/Resume&Results/Resume 3 - Krupakar Reddy P (1).xlsx"
    docx_path = "/home/great/claudeprojects/parser/parserdemo/Resume&Results/KrupakarReddy_SystemP.docx"

    print("=== KRUPAKAR REDDY RESUME VERIFICATION ANALYSIS ===\n")

    # Extract resume content
    print("1. EXTRACTING RESUME CONTENT...")
    resume_content = extract_docx_content(docx_path)

    if resume_content:
        print("Resume content extracted successfully")
        print(f"Resume length: {len(resume_content)} characters")
        print("\n--- RESUME CONTENT ---")
        print(resume_content)
        print("\n" + "="*80 + "\n")
    else:
        print("Failed to extract resume content")
        return

    # Extract verification data
    print("2. EXTRACTING VERIFICATION DATA...")
    verification_data = extract_excel_content(excel_path)

    if verification_data:
        print("Verification data extracted successfully")
        print(f"Sheets found: {list(verification_data.keys())}")

        for sheet_name, data in verification_data.items():
            print(f"\n--- SHEET: {sheet_name} ---")
            print(f"Rows: {len(data)}")

            if data:
                print("Sample data structure:")
                print(json.dumps(data[0], indent=2, default=str))

                # Print all data
                for i, row in enumerate(data):
                    print(f"\nRow {i+1}:")
                    for key, value in row.items():
                        if pd.notna(value):
                            print(f"  {key}: {value}")

        print("\n" + "="*80 + "\n")
    else:
        print("Failed to extract verification data")
        return

    # Analysis
    print("3. ANALYSIS AND FINDINGS...")

    # Save extracted data for further analysis
    analysis_result = {
        "resume_content": resume_content,
        "verification_data": verification_data,
        "analysis_timestamp": pd.Timestamp.now().isoformat()
    }

    output_path = "/home/great/claudeprojects/parser/parserdemo/krupakar_analysis_output.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(analysis_result, f, indent=2, default=str, ensure_ascii=False)

    print(f"Analysis data saved to: {output_path}")

    return analysis_result

if __name__ == "__main__":
    analyze_verification_data()