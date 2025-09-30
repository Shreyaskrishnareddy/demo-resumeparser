#!/usr/bin/env python3
"""
Comprehensive Validation Report
Compares parser output with actual resume content and verification Excel
"""

from fixed_comprehensive_parser import FixedComprehensiveParser
import fitz
from docx import Document
import json
import openpyxl
from datetime import datetime

def validate_resume(resume_name, text, filename, verification_data):
    """Validate a single resume against verification data"""
    parser = FixedComprehensiveParser()
    result = parser.parse_resume(text, filename)

    validation_results = {
        'resume_name': resume_name,
        'total_fields_checked': 0,
        'fields_present': 0,
        'fields_missing': 0,
        'issues_fixed': [],
        'issues_remaining': [],
        'new_fields_added': [],
        'field_details': {}
    }

    # Map verification fields to actual result
    for field_data in verification_data:
        category = field_data['category']
        field = field_data['field']
        expected = field_data.get('expected', '')

        validation_results['total_fields_checked'] += 1

        # Check actual value
        actual_value = check_field_value(result, category, field)
        is_present = actual_value is not None and str(actual_value).strip() != ''

        validation_results['field_details'][f"{category} - {field}"] = {
            'expected': expected,
            'actual': 'YES' if is_present else 'NO',
            'value': str(actual_value)[:100] if actual_value else 'MISSING'
        }

        if is_present:
            validation_results['fields_present'] += 1
            if expected and str(expected).upper() == 'NO':
                validation_results['issues_fixed'].append({
                    'category': category,
                    'field': field,
                    'value': str(actual_value)[:100]
                })
        else:
            validation_results['fields_missing'] += 1
            if expected and str(expected).upper() == 'YES':
                validation_results['issues_remaining'].append({
                    'category': category,
                    'field': field
                })

    # Add parsed data summary
    validation_results['parsed_data'] = {
        'work_experience_count': len(result.get('ListOfExperiences', [])),
        'skills_count': len(result.get('ListOfSkills', [])),
        'education_count': len(result.get('Education', [])),
        'certifications_count': len(result.get('Certifications', [])),
        'languages_count': len(result.get('Languages', [])),
        'domains_count': len(result.get('Domain', [])),
        'current_job_role': result.get('OverallSummary', {}).get('CurrentJobRole', 'MISSING'),
        'total_experience': result.get('OverallSummary', {}).get('TotalExperience', 'MISSING')
    }

    return validation_results, result

def check_field_value(result, category, field):
    """Extract actual field value from parsed result"""
    field_lower = field.lower()

    # Personal Details
    if category == 'Personal Details':
        personal = result.get('PersonalDetails', {})
        if 'full name' in field_lower:
            return personal.get('FullName')
        elif 'first name' in field_lower:
            return personal.get('FirstName')
        elif 'middle name' in field_lower:
            return personal.get('MiddleName')
        elif 'last name' in field_lower:
            return personal.get('LastName')
        elif 'email' in field_lower:
            return personal.get('EmailID')
        elif 'phone' in field_lower:
            return personal.get('PhoneNumber')
        elif 'country code' in field_lower:
            return personal.get('CountryCode')
        elif 'social media' in field_lower:
            return result.get('SocialMedia')

    # Overall Summary
    elif category == 'Overall Summary':
        summary = result.get('OverallSummary', {})
        if 'current job role' in field_lower:
            return summary.get('CurrentJobRole')
        elif 'relevant job titles' in field_lower:
            titles = summary.get('RelevantJobTitles', [])
            return titles if titles else None
        elif 'total experience' in field_lower:
            return summary.get('TotalExperience')
        elif 'summary' in field_lower or 'overall summary' in field_lower:
            return summary.get('OverallSummary')

    # Work Experiences
    elif category == 'Work Experiences':
        experiences = result.get('ListOfExperiences', [])
        if not experiences:
            return None
        if 'job title' in field_lower:
            return experiences[0].get('JobTitle') if experiences else None
        elif 'company' in field_lower:
            return experiences[0].get('CompanyName') if experiences else None
        elif 'location' in field_lower:
            return experiences[0].get('Location') if experiences else None
        elif 'start date' in field_lower:
            return experiences[0].get('StartDate') if experiences else None
        elif 'end date' in field_lower:
            return experiences[0].get('EndDate') if experiences else None
        elif 'summary' in field_lower or 'description' in field_lower:
            return experiences[0].get('Summary') if experiences else None
        elif 'employment type' in field_lower:
            return experiences[0].get('EmploymentType') if experiences else None
        elif 'total experience' in field_lower:
            return len(experiences) if experiences else None

    # Skills
    elif category == 'Skills':
        skills = result.get('ListOfSkills', [])
        if not skills:
            return None
        if 'skill' in field_lower and 'last used' not in field_lower:
            return [s.get('SkillName') for s in skills]
        elif 'last used' in field_lower:
            # Check if any skill has LastUsed field
            last_used_values = [s.get('LastUsed') for s in skills if s.get('LastUsed')]
            return last_used_values if last_used_values else None

    # Domain
    elif category == 'Domain':
        return result.get('Domain')

    # Education
    elif category == 'Education':
        education = result.get('Education', [])
        if not education:
            return None
        if 'degree' in field_lower or 'education detail' in field_lower:
            return education[0].get('Degree') if education else None
        elif 'institution' in field_lower or 'university' in field_lower:
            return education[0].get('Institution') if education else None
        elif 'field' in field_lower or 'major' in field_lower:
            return education[0].get('FieldOfStudy') if education else None
        elif 'location' in field_lower:
            return education[0].get('Location') if education else None

    # Certifications
    elif category == 'Certifications':
        certs = result.get('Certifications', [])
        if not certs:
            return None
        if 'name' in field_lower:
            return certs[0].get('Name') if certs else None
        elif 'issuer' in field_lower:
            return certs[0].get('Issuer') if certs else None
        elif 'year' in field_lower or 'date' in field_lower:
            return certs[0].get('IssuedDate') if certs else None

    # Languages
    elif category == 'Languages':
        languages = result.get('Languages', [])
        if languages:
            return [lang.get('Language') for lang in languages]

    # Projects
    elif category == 'Projects':
        projects = result.get('Projects', [])
        return projects if projects else None

    # Achievements
    elif category == 'Achievements':
        achievements = result.get('Achievements', [])
        return achievements if achievements else None

    # Key Responsibilities
    elif category == 'Key Responsibilities':
        experiences = result.get('ListOfExperiences', [])
        if experiences:
            summaries = [exp.get('Summary') for exp in experiences if exp.get('Summary')]
            return summaries if summaries else None

    return None

def main():
    print("="*100)
    print("🔬 COMPREHENSIVE RESUME PARSER VALIDATION REPORT")
    print("="*100)
    print(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()

    # Load verification Excel
    wb = openpyxl.load_workbook('Resume&Results/Parser Verification Results  (1).xlsx')
    sheet = wb['Sheet1']

    # Extract verification data
    verification_data = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row[0]:
            verification_data.append({
                'category': row[1],
                'field': row[2],
                'resume1': row[3],
                'resume2': row[4],
                'resume3': row[5]
            })

    all_results = {}
    all_parsed_json = {}

    # Resume 1: Ahmad Qasem
    print("\n" + "="*100)
    print("RESUME 1: AHMAD QASEM (PDF)")
    print("="*100)
    pdf = fitz.open('Resume&Results/Ahmad Qasem-Resume.pdf')
    text = ''.join([page.get_text() for page in pdf])
    resume1_verification = [{**v, 'expected': v['resume1']} for v in verification_data]
    result1, parsed1 = validate_resume('Ahmad Qasem', text, 'Ahmad Qasem-Resume.pdf', resume1_verification)
    all_results['Resume 1 - Ahmad Qasem'] = result1
    all_parsed_json['Resume 1 - Ahmad Qasem'] = parsed1

    print(f"\nStatistics:")
    print(f"  Total fields checked: {result1['total_fields_checked']}")
    print(f"  Fields present: {result1['fields_present']} ({result1['fields_present']/result1['total_fields_checked']*100:.1f}%)")
    print(f"  Fields missing: {result1['fields_missing']} ({result1['fields_missing']/result1['total_fields_checked']*100:.1f}%)")
    print(f"\n  Issues fixed: {len(result1['issues_fixed'])}")
    for fix in result1['issues_fixed'][:5]:
        print(f"     - {fix['category']}: {fix['field']}")
    print(f"\n  Issues remaining: {len(result1['issues_remaining'])}")
    for issue in result1['issues_remaining'][:5]:
        print(f"     - {issue['category']}: {issue['field']}")

    print(f"\nParsed Data:")
    for key, value in result1['parsed_data'].items():
        print(f"  {key}: {value}")

    # Resume 2: Zamen Aladwani
    print("\n" + "="*100)
    print("RESUME 2: ZAMEN ALADWANI (PDF)")
    print("="*100)
    pdf = fitz.open('Resume&Results/ZAMEN_ALADWANI_PROJECT MANAGER_09_01_2023.pdf')
    text = ''.join([page.get_text() for page in pdf])
    resume2_verification = [{**v, 'expected': v['resume2']} for v in verification_data]
    result2, parsed2 = validate_resume('Zamen Aladwani', text, 'ZAMEN_ALADWANI_PROJECT MANAGER_09_01_2023.pdf', resume2_verification)
    all_results['Resume 2 - Zamen Aladwani'] = result2
    all_parsed_json['Resume 2 - Zamen Aladwani'] = parsed2

    print(f"\nStatistics:")
    print(f"  Total fields checked: {result2['total_fields_checked']}")
    print(f"  Fields present: {result2['fields_present']} ({result2['fields_present']/result2['total_fields_checked']*100:.1f}%)")
    print(f"  Fields missing: {result2['fields_missing']} ({result2['fields_missing']/result2['total_fields_checked']*100:.1f}%)")
    print(f"\n  Issues fixed: {len(result2['issues_fixed'])}")
    print(f"\n  Issues remaining: {len(result2['issues_remaining'])}")

    print(f"\nParsed Data:")
    for key, value in result2['parsed_data'].items():
        print(f"  {key}: {value}")

    # Resume 3: Krupakar Reddy
    print("\n" + "="*100)
    print("RESUME 3: KRUPAKAR REDDY (DOCX)")
    print("="*100)
    doc = Document('Resume&Results/KrupakarReddy_SystemP.docx')
    text = '\n'.join([p.text for p in doc.paragraphs])
    resume3_verification = [{**v, 'expected': v['resume3']} for v in verification_data]
    result3, parsed3 = validate_resume('Krupakar Reddy', text, 'KrupakarReddy_SystemP.docx', resume3_verification)
    all_results['Resume 3 - Krupakar Reddy'] = result3
    all_parsed_json['Resume 3 - Krupakar Reddy'] = parsed3

    print(f"\nStatistics:")
    print(f"  Total fields checked: {result3['total_fields_checked']}")
    print(f"  Fields present: {result3['fields_present']} ({result3['fields_present']/result3['total_fields_checked']*100:.1f}%)")
    print(f"  Fields missing: {result3['fields_missing']} ({result3['fields_missing']/result3['total_fields_checked']*100:.1f}%)")
    print(f"\n  Issues fixed: {len(result3['issues_fixed'])}")
    print(f"\n  Issues remaining: {len(result3['issues_remaining'])}")

    print(f"\nParsed Data:")
    for key, value in result3['parsed_data'].items():
        print(f"  {key}: {value}")

    # Resume 4: Venkat Rohit (not in verification Excel but we'll validate anyway)
    print("\n" + "="*100)
    print("RESUME 4: VENKAT ROHIT (DOCX)")
    print("="*100)
    doc = Document('Resume&Results/Venkat_Rohit_Senior .NET Full Stack Developer (1).docx')
    text = '\n'.join([p.text for p in doc.paragraphs])
    # Use Resume 3 verification as template
    resume4_verification = [{**v, 'expected': ''} for v in verification_data]
    result4, parsed4 = validate_resume('Venkat Rohit', text, 'Venkat_Rohit_Senior .NET Full Stack Developer (1).docx', resume4_verification)
    all_results['Resume 4 - Venkat Rohit'] = result4
    all_parsed_json['Resume 4 - Venkat Rohit'] = parsed4

    print(f"\nStatistics:")
    print(f"  Total fields checked: {result4['total_fields_checked']}")
    print(f"  Fields present: {result4['fields_present']} ({result4['fields_present']/result4['total_fields_checked']*100:.1f}%)")
    print(f"  Fields missing: {result4['fields_missing']} ({result4['fields_missing']/result4['total_fields_checked']*100:.1f}%)")

    print(f"\nParsed Data:")
    for key, value in result4['parsed_data'].items():
        print(f"  {key}: {value}")

    # Overall Summary
    print("\n" + "="*100)
    print("OVERALL SUMMARY")
    print("="*100)
    total_fixed = sum(len(r['issues_fixed']) for r in all_results.values())
    total_remaining = sum(len(r['issues_remaining']) for r in all_results.values())
    total_fields_present = sum(r['fields_present'] for r in all_results.values())
    total_fields_checked = sum(r['total_fields_checked'] for r in all_results.values())

    print(f"\nTotal issues FIXED across all resumes: {total_fixed}")
    print(f"Total issues REMAINING: {total_remaining}")
    print(f"Overall field coverage: {total_fields_present}/{total_fields_checked} ({total_fields_present/total_fields_checked*100:.1f}%)")

    # Save results
    with open('validation_results.json', 'w', encoding='utf-8') as f:
        json.dump(all_results, f, indent=2, ensure_ascii=False)

    with open('all_resumes_parsed.json', 'w', encoding='utf-8') as f:
        json.dump(all_parsed_json, f, indent=2, ensure_ascii=False, default=str)

    print(f"\nResults saved to:")
    print(f"  - validation_results.json")
    print(f"  - all_resumes_parsed.json")
    print("\n" + "="*100)

if __name__ == '__main__':
    main()