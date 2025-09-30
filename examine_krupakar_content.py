from docx import Document

# Examine the actual content structure of Krupakar's resume
doc = Document('Resume&Results/KrupakarReddy_SystemP.docx')

print("="*80)
print("📄 KRUPAKAR RESUME - RAW CONTENT ANALYSIS")
print("="*80)

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# Look for education section
print("\n🔍 SEARCHING FOR EDUCATION SECTION:")
education_keywords = ['education', 'academic', 'qualification', 'degree', 'university', 'college', 'bachelor', 'master']

education_found = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if any(keyword in text.lower() for keyword in education_keywords):
        print(f"\nLine {i}: {repr(text)}")
        # Show context (5 lines before and after)
        start = max(0, i-2)
        end = min(len(doc.paragraphs), i+10)
        print(f"\n   Context (lines {start}-{end}):")
        for j in range(start, end):
            marker = ">>>" if j == i else "   "
            print(f"{marker} {j:3d}: {repr(doc.paragraphs[j].text.strip())}")
        education_found = True
        print("\n" + "-"*80)

if not education_found:
    print("   ❌ No obvious education section headers found")

# Look for work experience
print("\n🔍 SEARCHING FOR WORK EXPERIENCE:")
experience_keywords = ['experience', 'employment', 'work history', 'professional', 'career']

experience_found = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if any(keyword in text.lower() for keyword in experience_keywords) and len(text) < 50:
        print(f"\nLine {i}: {repr(text)}")
        # Show context
        start = max(0, i-2)
        end = min(len(doc.paragraphs), i+15)
        print(f"\n   Context (lines {start}-{end}):")
        for j in range(start, end):
            marker = ">>>" if j == i else "   "
            print(f"{marker} {j:3d}: {repr(doc.paragraphs[j].text.strip())}")
        experience_found = True
        break

# Show all content to understand structure
print("\n\n📝 FULL RESUME CONTENT (first 100 lines):")
print("="*80)
for i, para in enumerate(doc.paragraphs[:100]):
    text = para.text.strip()
    if text:  # Only show non-empty lines
        print(f"{i:3d}: {text}")

print("\n" + "="*80)
print(f"Total non-empty lines: {sum(1 for p in doc.paragraphs if p.text.strip())}")
print("="*80)